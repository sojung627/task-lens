import base64
import binascii
import io
from pathlib import Path
from uuid import uuid4

from backend.app.core.config import Settings
from backend.app.schemas.chat import FileUploadPayload


class FileValidationError(Exception):
    pass


class FileReadError(Exception):
    pass


class FileSizeLimitError(FileValidationError):
    pass


TEXT_EXTENSIONS = {
    "txt",
    "md",
    "csv",
    "json",
    "xml",
    "yaml",
    "yml",
    "py",
    "js",
    "jsx",
    "ts",
    "tsx",
    "html",
    "css",
    "sql",
    "java",
    "c",
    "cpp",
    "h",
    "log",
}
DOCUMENT_EXTENSIONS = TEXT_EXTENSIONS | {"pdf", "docx"}
AUDIO_EXTENSIONS = {"mp3", "mp4", "mpeg", "mpga", "m4a", "ogg", "wav", "webm", "flac"}
UPLOAD_EXTENSIONS = DOCUMENT_EXTENSIONS | AUDIO_EXTENSIONS
GENERATED_EXTENSIONS = TEXT_EXTENSIONS | {"pdf", "docx"}


class FileService:
    def __init__(self, settings: Settings):
        self.settings = settings
        self.upload_directory = settings.storage_directory / "files"
        self.upload_directory.mkdir(parents=True, exist_ok=True)

    def decode_upload(self, upload: FileUploadPayload) -> tuple[bytes, str, str]:
        safe_name = Path(upload.name).name
        if safe_name != upload.name or not safe_name:
            raise FileValidationError("올바르지 않은 파일명이에요.")

        extension = Path(safe_name).suffix.lower().lstrip(".")
        if extension not in UPLOAD_EXTENSIONS:
            supported = ", ".join(sorted(UPLOAD_EXTENSIONS))
            raise FileValidationError(
                f"{safe_name} 파일 형식은 지원하지 않아요. 지원 형식: {supported}"
            )

        try:
            content = base64.b64decode(upload.content_base64, validate=True)
        except (binascii.Error, ValueError) as exc:
            raise FileValidationError(f"{safe_name} 파일 데이터가 올바르지 않아요.") from exc

        if not content:
            raise FileValidationError(f"{safe_name} 파일이 비어 있어요.")
        if len(content) > self.settings.max_chat_file_bytes:
            limit_in_kilobytes = self.settings.max_chat_file_bytes / 1024
            limit_label = (
                f"{int(limit_in_kilobytes)}KB"
                if limit_in_kilobytes.is_integer()
                else f"{limit_in_kilobytes:.1f}KB"
            )
            raise FileSizeLimitError(
                f"{safe_name} 파일은 {limit_label} 이하만 첨부할 수 있어요. "
                "더 큰 파일은 내용을 나누어 업로드해 주세요."
            )
        if len(content) > self.settings.max_upload_bytes:
            raise FileValidationError(f"{safe_name} 파일은 업로드 제한 용량을 초과했어요.")
        return content, safe_name, extension

    def validate_generated_file(self, name: str, content: str) -> tuple[str, str, bytes]:
        safe_name = Path(name).name
        if safe_name != name or not safe_name:
            raise FileValidationError("AI가 생성한 파일명이 올바르지 않아요.")
        extension = Path(safe_name).suffix.lower().lstrip(".")
        if extension not in GENERATED_EXTENSIONS:
            supported = ", ".join(sorted(ext.upper() for ext in GENERATED_EXTENSIONS))
            raise FileValidationError(f"생성 파일은 {supported} 형식만 지원해요.")

        if extension == "pdf":
            encoded = self._render_pdf(content)
        elif extension == "docx":
            encoded = self._render_docx(content)
        else:
            encoded = content.encode("utf-8")
        if len(encoded) > self.settings.max_generated_file_bytes:
            raise FileValidationError("AI가 생성한 파일이 허용 용량을 초과했어요.")
        return safe_name, extension, encoded


    @property
    def generated_extensions(self) -> set[str]:
        return set(GENERATED_EXTENSIONS)

    @staticmethod
    def mime_type_for_extension(extension: str) -> str:
        mime_types = {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "json": "application/json",
            "xml": "application/xml",
            "html": "text/html",
            "css": "text/css",
            "csv": "text/csv",
        }
        return mime_types.get(extension, "text/plain")

    @staticmethod
    def _render_docx(content: str) -> bytes:
        from docx import Document

        buffer = io.BytesIO()
        document = Document()
        for line in content.splitlines() or [content]:
            document.add_paragraph(line)
        document.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _render_pdf(content: str) -> bytes:
        """AI가 만든 UTF-8 텍스트를 실제 PDF 바이트로 변환한다."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        buffer = io.BytesIO()
        pdfmetrics.registerFont(UnicodeCIDFont("HYSMyeongJo-Medium"))
        document = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=20 * mm,
            rightMargin=20 * mm,
            topMargin=20 * mm,
            bottomMargin=20 * mm,
            title="TaskLens generated document",
        )
        style = ParagraphStyle(
            name="TaskLensKorean",
            fontName="HYSMyeongJo-Medium",
            fontSize=10.5,
            leading=16,
            wordWrap="CJK",
        )
        story = []
        for line in content.splitlines() or [content]:
            escaped = (
                line.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            story.append(Paragraph(escaped or " ", style))
            story.append(Spacer(1, 2 * mm))
        document.build(story)
        return buffer.getvalue()

    def save(self, content: bytes, extension: str) -> str:
        stored_name = f"{uuid4().hex}.{extension}"
        target = self.upload_directory / stored_name
        target.write_bytes(content)
        return stored_name

    def delete_stored(self, stored_name: str) -> None:
        path = self.upload_directory / Path(stored_name).name
        if path.is_file():
            path.unlink()

    def read_stored(self, stored_name: str) -> bytes:
        path = self.upload_directory / Path(stored_name).name
        if not path.is_file():
            raise FileReadError("파일을 찾을 수 없어요.")
        return path.read_bytes()

    def extract_text(self, content: bytes, extension: str) -> str:
        try:
            if extension in TEXT_EXTENSIONS:
                text = content.decode("utf-8-sig")
            elif extension == "pdf":
                from pypdf import PdfReader

                reader = PdfReader(io.BytesIO(content))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
            elif extension == "docx":
                from docx import Document

                document = Document(io.BytesIO(content))
                text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            elif extension in AUDIO_EXTENSIONS:
                return "[음성 파일은 음성 인식 단계에서 텍스트로 변환됩니다.]"
            else:
                raise FileValidationError("지원하지 않는 파일 형식이에요.")
        except UnicodeDecodeError as exc:
            raise FileReadError(f"{extension.upper()} 파일의 문자 인코딩을 읽지 못했어요.") from exc
        except FileValidationError:
            raise
        except Exception as exc:
            raise FileReadError(f"{extension.upper()} 파일을 읽지 못했어요.") from exc

        normalized = text.strip()
        if not normalized:
            return "[파일에 읽을 수 있는 텍스트가 없음]"
        if len(normalized) > self.settings.max_file_text_length:
            return normalized[: self.settings.max_file_text_length] + "\n[이후 내용은 길이 제한으로 생략됨]"
        return normalized